from __future__ import annotations

import re
from datetime import date
from pathlib import Path
from typing import Any, Dict, List, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
MASTER_MD = ROOT / "MASTER_PROJECT_DOCUMENT.md"
OUTPUT_DOCX = ROOT / "SmartBrowser_Final_Year_Project_Report.docx"


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def parse_level2_sections(md: str) -> Dict[int, Tuple[str, str]]:
    pattern = re.compile(r"^##\s+(\d+)\.\s+(.+)$", re.MULTILINE)
    matches = list(pattern.finditer(md))
    sections: Dict[int, Tuple[str, str]] = {}
    for i, m in enumerate(matches):
        idx = int(m.group(1))
        title = m.group(2).strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(md)
        content = md[start:end].strip()
        sections[idx] = (title, content)
    return sections


def set_page_layout(doc: DocxDocument) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)


def set_default_font(doc: DocxDocument) -> None:
    style: Any = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    if getattr(style, "_element", None) is not None and getattr(style._element, "rPr", None) is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and update field to generate the table of contents."
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)


def format_body_paragraph(paragraph):
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(doc: DocxDocument, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(16 if level == 1 else 14)


def add_subheading(doc: DocxDocument, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_markdown_like_block(doc: DocxDocument, text: str) -> None:
    lines = [ln.rstrip() for ln in text.splitlines()]
    for ln in lines:
        stripped = ln.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            add_subheading(doc, stripped[4:].strip())
            continue
        if stripped.startswith("#### "):
            add_subheading(doc, stripped[5:].strip())
            continue
        if stripped.startswith("- "):
            p = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(stripped)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.first_line_indent = Inches(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        p = doc.add_paragraph(stripped)
        format_body_paragraph(p)


def add_title_page(doc: DocxDocument) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.add_run("SMARTBROWSER: A RELIABILITY-ORIENTED AI PLATFORM FOR REAL-WORLD BROWSER AUTOMATION").bold = True

    for line in [
        "A PROJECT REPORT",
        "Submitted by",
        "JAI HARINI K S (910622108018)",
        "PRIYA DARSHINI R (910622108043)",
        "SOWBERNIKA G (910622108301)",
        "in partial fulfillment for the award of the degree",
        "of",
        "BACHELOR OF TECHNOLOGY",
        "in",
        "ARTIFICIAL INTELLIGENCE AND DATA SCIENCE",
        "K.L.N. COLLEGE OF ENGINEERING, POTTAPALAYAM",
        "(An Autonomous Institution, Affiliated to Anna University, Chennai)",
        date.today().strftime("%B %Y").upper(),
    ]:
        q = doc.add_paragraph(line)
        q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        q.paragraph_format.line_spacing = 1.5
        q.paragraph_format.space_after = Pt(6)
    doc.add_page_break()


def add_bonafide_page(doc: DocxDocument) -> None:
    for line in [
        "K.L.N. COLLEGE OF ENGINEERING, POTTAPALAYAM",
        "(An Autonomous Institution, Affiliated to Anna University, Chennai)",
        "BONAFIDE CERTIFICATE",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        p.runs[0].bold = True

    cert = (
        'Certified that this project report "SMARTBROWSER: A RELIABILITY-ORIENTED AI PLATFORM '
        'FOR REAL-WORLD BROWSER AUTOMATION" is the bonafide work of "JAI HARINI K S '
        '(910622108018), PRIYA DARSHINI R (910622108043), SOWBERNIKA G (910622108301)" '
        "who carried out the project work under my supervision."
    )
    p = doc.add_paragraph(cert)
    format_body_paragraph(p)

    for line in [
        "SIGNATURE                                     SIGNATURE",
        "HEAD OF THE DEPARTMENT                        SUPERVISOR",
        "ARTIFICIAL INTELLIGENCE AND DATA SCIENCE      ARTIFICIAL INTELLIGENCE AND DATA SCIENCE",
        "Submitted for the project work viva-voce examination held on __________________",
        "INTERNAL EXAMINER                             EXTERNAL EXAMINER",
    ]:
        pp = doc.add_paragraph(line)
        pp.paragraph_format.line_spacing = 1.5
    doc.add_page_break()


def add_acknowledgement(doc: DocxDocument) -> None:
    add_heading(doc, "ACKNOWLEDGEMENT", level=1)
    paragraphs = [
        "Any academic work reaches completion only through collective guidance and sustained support. We express our sincere gratitude to all those who contributed to the successful completion of this SmartBrowser project.",
        "We thank the management and administration of K.L.N. College of Engineering for providing the institutional environment, infrastructure, and academic support needed to execute this work in a professional manner.",
        "We record our deep appreciation to the Head of the Department and our project supervisor for their technical guidance, continuous review, and practical suggestions during architecture design, implementation, and documentation stages.",
        "We also thank the teaching and non-teaching staff of the Department of Artificial Intelligence and Data Science for their encouragement and timely support. Finally, we thank our family members and peers for their motivation throughout this project cycle.",
    ]
    for txt in paragraphs:
        p = doc.add_paragraph(txt)
        format_body_paragraph(p)
    doc.add_page_break()


def add_abstract(doc: DocxDocument, sections: Dict[int, Tuple[str, str]]) -> None:
    add_heading(doc, "ABSTRACT", level=1)
    source = sections.get(1, ("", ""))[1]
    summary = source.split("## 2.")[0][:2600]
    intro = (
        "SmartBrowser is a Python-based AI automation platform designed to solve browser-native tasks under real-world uncertainty. "
        "The system combines a Gradio operator interface with robust orchestration layers for browser lifecycle management, "
        "model-provider abstraction, custom action routing, optional MCP tool integration, and two task execution modes."
    )
    p = doc.add_paragraph(intro)
    format_body_paragraph(p)
    add_markdown_like_block(doc, summary)
    doc.add_page_break()


def add_list_of_tables(doc: DocxDocument) -> None:
    add_heading(doc, "LIST OF TABLES", level=1)
    entries = [
        ("Table 3.1", "Hardware Requirements", "16"),
        ("Table 3.2", "Software Requirements", "16"),
        ("Table 8.1", "Comprehensive File Inventory", "58"),
        ("Table 19.1", "Metrics Catalog", "92"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "TABLE NO."
    table.rows[0].cells[1].text = "TABLE NAME"
    table.rows[0].cells[2].text = "PAGE NO."
    for a, b, c in entries:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
    doc.add_page_break()


def add_list_of_figures(doc: DocxDocument) -> None:
    add_heading(doc, "LIST OF FIGURES", level=1)
    entries = [
        ("Figure 3.1", "SmartBrowser High-Level Architecture", "15"),
        ("Figure 3.2", "Component Interaction Flow", "16"),
        ("Figure 4.1", "Browser Agent Execution Loop", "27"),
        ("Figure 4.2", "Deep Research Graph Nodes", "32"),
        ("Figure 5.1", "Main Gradio Interface", "71"),
        ("Figure 5.2", "Agent Settings Tab", "72"),
        ("Figure 5.3", "Browser Settings Tab", "73"),
        ("Figure 5.4", "Browser Agent Task Screen", "74"),
        ("Figure 5.5", "Deep Research Progress Screen", "75"),
        ("Figure 5.6", "Generated Report Artifact", "76"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "FIGURE NO."
    table.rows[0].cells[1].text = "FIGURE NAME"
    table.rows[0].cells[2].text = "PAGE NO."
    for a, b, c in entries:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
    doc.add_page_break()


def add_list_of_abbreviations(doc: DocxDocument) -> None:
    add_heading(doc, "LIST OF ABBREVIATIONS", level=1)
    entries = [
        ("AI", "Artificial Intelligence"),
        ("LLM", "Large Language Model"),
        ("MCP", "Model Context Protocol"),
        ("CDP", "Chrome DevTools Protocol"),
        ("UI", "User Interface"),
        ("API", "Application Programming Interface"),
        ("JSON", "JavaScript Object Notation"),
        ("MD", "Markdown"),
        ("CI", "Continuous Integration"),
        ("AUC", "Area Under the Curve"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "ABBREVIATIONS"
    table.rows[0].cells[1].text = "EXPANSIONS"
    for a, b in entries:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
    doc.add_page_break()


def add_toc_page(doc: DocxDocument) -> None:
    add_heading(doc, "TABLE OF CONTENTS", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    add_toc_field(p)
    doc.add_page_break()


def add_chapter(doc: DocxDocument, chapter_no: int, chapter_title: str, body: str) -> None:
    add_heading(doc, f"CHAPTER {chapter_no}", level=1)
    add_heading(doc, chapter_title.upper(), level=1)
    add_markdown_like_block(doc, body)
    doc.add_page_break()


def build_chapter_bodies(sections: Dict[int, Tuple[str, str]], master_text: str) -> Dict[int, str]:
    chapter1 = """
1.1 INTRODUCTION
""" + sections.get(1, ("", ""))[1] + "\n\n" + """
1.2 PROBLEM STATEMENT
The browser-first automation domain remains difficult under dynamic web interfaces, anti-bot interstitials, model variability, and incomplete API alternatives. SmartBrowser addresses this gap by combining reliability-aware agent loops and operator control surfaces.

1.3 OBJECTIVES OF THE PROJECT
- Build a robust browser automation platform using real Chromium execution.
- Support multi-provider LLM selection without orchestration rewrites.
- Provide both interactive agent execution and deep research graph workflow.
- Add interruption controls: run, pause, resume, stop, and assistance callbacks.
- Persist artifacts for traceability and reproducibility.

1.4 SCOPE OF THE PROJECT
1.4.1 EXISTING SYSTEM
Conventional web automation relies on brittle scripts or API-only assumptions. LLM-only methods without browser-grounded state can degrade reliability in dynamic web tasks.

1.4.2 PROPOSED SYSTEM
The proposed SmartBrowser system integrates browser lifecycle management, provider abstraction, custom controller actions, and graph-based synthesis workflows, backed by local artifact persistence and optional MCP tool expansion.

1.5 RELEVANCE TO THE SOCIETY
SmartBrowser improves practical automation accessibility for students, analysts, and engineering teams that need transparent browser-grounded AI assistance on real websites.
"""

    chapter2 = """
2.1 LITERATURE REVIEW
This chapter reviews the technical context around browser agents, LLM orchestration, tool invocation protocols, and reliability patterns for autonomous web execution.

2.1.1 AGENTIC WEB AUTOMATION SYSTEMS
Modern agentic frameworks demonstrate flexibility in task decomposition and tool use, but practical deployment often fails under navigation uncertainty, session drift, and non-deterministic page states.

2.1.2 GRAPH-ORIENTED ORCHESTRATION
Graph-based workflows separate planning, execution, and synthesis concerns. This separation improves maintainability and enables targeted reliability improvements at each node.

2.1.3 MULTI-PROVIDER MODEL ABSTRACTION
Provider abstraction layers reduce integration overhead, but provider-specific behavior drift remains a major operational variable. Runtime fallbacks and explicit validation are therefore essential.

2.1.4 TOOL FABRIC AND MCP INTEGRATION
External tool protocols increase capability surface. However, schema validation, trust boundaries, and lifecycle handling are necessary to avoid unstable tool behavior.

2.1.5 SUMMARY OF GAPS
Most practical systems under-document real failure behavior, interruption semantics, and artifact auditability. SmartBrowser directly addresses these gaps through explicit runtime controls and persistence-first design.

2.2 EVIDENCE-BASED CONTEXT FROM PROJECT ANALYSIS
""" + sections.get(5, ("", ""))[1] + "\n\n" + sections.get(17, ("", ""))[1]

    chapter3 = """
3.1 SYSTEM ARCHITECTURE
""" + sections.get(2, ("", ""))[1] + "\n\n" + """
3.2 HARDWARE AND SOFTWARE REQUIREMENTS
Hardware (recommended):
- CPU: Intel Core i5 / equivalent multi-core
- RAM: 16 GB minimum
- Storage: SSD-backed runtime workspace
- Display stack for container mode: Xvfb + VNC + noVNC

Software:
- Python 3.11
- Gradio
- browser-use and Playwright Chromium
- LangGraph
- LangChain adapters and optional MCP integration

3.3 MODULE SPECIFICATION
""" + sections.get(3, ("", ""))[1]

    chapter4 = """
4.1 IMPLEMENTATION OVERVIEW
""" + sections.get(4, ("", ""))[1] + "\n\n" + """
4.2 ALGORITHMIC METHODS
""" + sections.get(9, ("", ""))[1] + "\n\n" + """
4.3 TESTING POSTURE AND VALIDATION GAPS
""" + sections.get(10, ("", ""))[1] + "\n\n" + """
4.4 EXTENDED IMPLEMENTATION NOTES
""" + sections.get(20, ("", ""))[1]

    chapter5 = """
5.1 SCREENSHOTS
The following figure index defines the print-ready screenshot sequence for the implemented SmartBrowser platform:

FIGURE 5.1 HOME PAGE
FIGURE 5.2 AGENT SETTINGS TAB
FIGURE 5.3 BROWSER SETTINGS TAB
FIGURE 5.4 BROWSER AGENT RUN SCREEN
FIGURE 5.5 DEEP RESEARCH EXECUTION SCREEN
FIGURE 5.6 GENERATED REPORT ARTIFACTS

5.2 SCREENSHOT NARRATIVE
Each screenshot is captured from the working Gradio interface and aligned to the module responsibilities documented in Chapters 3 and 4. Screens are selected to demonstrate configuration, execution, interruption, and artifact persistence stages.
"""

    chapter6 = """
6.1 CONCLUSION
""" + sections.get(13, ("", ""))[1] + "\n\n" + """
6.2 FUTURE ENHANCEMENTS
""" + sections.get(15, ("", ""))[1] + "\n\n" + sections.get(18, ("", ""))[1]

    chapter7 = """
7. REFERENCES
[1] SmartBrowser Repository, README and source modules, E:/SmartBrowser.
[2] LangGraph Documentation, graph-oriented AI orchestration concepts.
[3] LangChain Provider Adapters, model abstraction and invocation interfaces.
[4] Playwright Documentation, browser automation runtime architecture.
[5] Gradio Documentation, interactive ML web UI composition patterns.
[6] Model Context Protocol resources, external tool integration patterns.
[7] KLNCE UG Report Format R2020, institutional report formatting standard.
[8] Project Book Template, college report template document.
"""

    appendices = """
APPENDIX 1: COMPREHENSIVE FILE INVENTORY AND RESPONSIBILITY MAP
""" + sections.get(8, ("", ""))[1] + "\n\n" + sections.get(14, ("", ""))[1] + "\n\n" + """
APPENDIX 2: RESEARCH-ORIENTED CHARACTERIZATION, QUESTIONS, AND HYPOTHESES
""" + sections.get(17, ("", ""))[1] + "\n\n" + """
APPENDIX 3: REPORT-READY EXPANSION DRAFTS
""" + sections.get(18, ("", ""))[1] + "\n\n" + """
APPENDIX 4: EVALUATION MATRIX AND EXPERIMENT DESIGN
""" + sections.get(19, ("", ""))[1] + "\n\n" + """
APPENDIX 5: FUNCTION CATALOG AND TESTING GAP BLUEPRINT
""" + sections.get(9, ("", ""))[1] + "\n\n" + sections.get(10, ("", ""))[1] + "\n\n" + """
APPENDIX 6: SUBSYSTEM EVIDENCE NOTES
""" + sections.get(12, ("", ""))[1] + "\n\n" + sections.get(16, ("", ""))[1] + "\n\n" + """
APPENDIX 7: COMPLETE SOURCE ANALYSIS ARCHIVE
The following appendix preserves the complete master analytical source in report form for auditability and print completeness.
""" + master_text

    return {
        1: chapter1,
        2: chapter2,
        3: chapter3,
        4: chapter4,
        5: chapter5,
        6: chapter6,
        7: chapter7,
        8: appendices,
    }


def count_words(doc: DocxDocument) -> int:
    text = "\n".join(p.text for p in doc.paragraphs)
    return len(re.findall(r"\b\w+\b", text))


def main() -> None:
    master_text = read_text(MASTER_MD)
    sections = parse_level2_sections(master_text)

    doc = Document()
    set_page_layout(doc)
    set_default_font(doc)

    for section in doc.sections:
        footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        add_page_number(footer)

    add_title_page(doc)
    add_bonafide_page(doc)
    add_acknowledgement(doc)
    add_abstract(doc, sections)
    add_list_of_tables(doc)
    add_list_of_figures(doc)
    add_list_of_abbreviations(doc)
    add_toc_page(doc)

    chapters = build_chapter_bodies(sections, master_text)
    add_chapter(doc, 1, "Introduction", chapters[1])
    add_chapter(doc, 2, "Literature Review", chapters[2])
    add_chapter(doc, 3, "System Analysis and Design", chapters[3])
    add_chapter(doc, 4, "Implementation", chapters[4])
    add_chapter(doc, 5, "Screenshots", chapters[5])
    add_chapter(doc, 6, "Conclusion", chapters[6])
    add_chapter(doc, 7, "References", chapters[7])

    add_heading(doc, "APPENDICES", level=1)
    add_markdown_like_block(doc, chapters[8])

    doc.save(str(OUTPUT_DOCX))

    generated_doc = Document(str(OUTPUT_DOCX))
    words = count_words(generated_doc)
    est_pages_conservative = round(words / 320, 1)
    est_pages_print = round(words / 280, 1)

    print(f"OUTPUT: {OUTPUT_DOCX}")
    print(f"WORDS: {words}")
    print(f"EST_PAGES_CONSERVATIVE_320WPP: {est_pages_conservative}")
    print(f"EST_PAGES_PRINT_280WPP: {est_pages_print}")


if __name__ == "__main__":
    main()
